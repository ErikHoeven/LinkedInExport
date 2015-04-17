'''
Created on Mar 19, 2015

@author: erik
'''

from docx import Document
from getProfile import getFullName
from getProfile import getBackgroundSummary
import os
from getProfile import getBackgroundExperience
from getProfile import getBackgroundProjects
from docx.shared import Pt



def setProfile(profile):
     

    #Inialize variables
    os.chdir("/home/erik/workspace/LinkedInMongo/MongoRead/Export")
    i = 0
    documentName = ""
    
    def setFullName(FullName):
        
        FullNameLabel = table.cell(0,0)
        FullNameText  = table.cell(0,1)
        
        FullNameLabel.text = "Naam:"
        FullNameText.text  = FullName
        print FullName
               
    def setAdress(Adres):
        AdresLabel = table.cell(1,0)
        AdresText  = table.cell(1,1)
            
        AdresLabel.text = "Adres:"
        AdresText.text  = Adres
        
       
    def setPostcode(Postcode):
        PostCodeLabel = table.cell(2,0)
        PostCodeText  = table.cell(2,1)
            
        PostCodeLabel.text = "Postcode:"
        PostCodeText.text  = Postcode
       
        
    def setGSM(GSM):
        GSMLabel = table.cell(3,0)
        GSMText  = table.cell(3,1)
            
        GSMLabel.text = "GSM:"
        GSMText.text  = GSM
        
        
    def setEmail(Email):
        EmailLabel = table.cell(4,0)
        EmailText  = table.cell(4,1)
            
        EmailLabel.text = "Email:"
        EmailText.text  = Email
    
    def setGebDat(gebdat):
        gebdatLabel = table.cell(0,2)
        gebdatText  = table.cell(0,3)
            
        gebdatLabel.text = "Geboorte datum:"
        gebdatText.text  = gebdat
        
      
    def setGender(Gender):
        GenderLabel = table.cell(1,2)
        GenderText  = table.cell(1,3)
            
        GenderLabel.text = "Geslacht:"
        GenderText.text  = Gender
        
    def setNationality(Nationality):
        NationalityLabel = table.cell(1,2)
        NationalityText  = table.cell(1,3)
            
        NationalityLabel.text = "Nationaliteit:"
        NationalityText.text  = Nationality
               
    def setSummary(Summary ):
        print "Start setSummary"
        #print Summary
        counter = 0
        SummaryText = ""
        print len(Summary) 
        
        if len(Summary) > 0:
                    
                while (counter < len(Summary)): 
                    SummaryText = SummaryText + " " + Summary[counter]
                    counter = counter + 1
                print "Einde setSummary"
            
        if len(Summary) == 0:
                SummaryText = "Geen achtergrond beschikbaar"
        return SummaryText
        
        
        
    def setBackgroundExperience(BackgroundExperience ):
        print "Start BackgroundExperience"
        print type(BackgroundExperience)
        print isinstance(BackgroundExperience,unicode)
        if isinstance(BackgroundExperience,unicode) != True:
            print "type is geen string"
            counter = 0
            BackgroundExperienceText = ""     

            #print Summary
            BackgroundExperienceText = ""
           
            if len(BackgroundExperience) > 0:
                    
                while (counter < len(BackgroundExperience)): 
                    BackgroundExperienceText = BackgroundExperienceText + " " + BackgroundExperience[counter]
                    counter = counter + 1
                print "Einde setBackgroundExperience"
            
            if len(BackgroundExperienceText) == 0:
                    BackgroundExperienceText = "Geen ervaring beschikbaar"
        
        
        if isinstance(BackgroundExperience,unicode) == True:
            print "type is string"
            BackgroundExperienceText = BackgroundExperience
            
        return BackgroundExperienceText    

    def setProjectExperience(ProjectExperience ):
        print "Start ProjectExperience"
        print type(ProjectExperience)
        print isinstance(ProjectExperience,unicode)
        if isinstance(ProjectExperience,unicode) != True:
            print "type is geen string"
            counter = 0   
            #print Summary
            ProjectExperienceText = ""
           
            if len(ProjectExperience) > 0:
                    
                while (counter < len(ProjectExperience)): 
                    ProjectExperienceText = ProjectExperienceText + " " + ProjectExperience[counter]
                    counter = counter + 1
                print "Einde setBackgroundExperience"
            
            if len(ProjectExperienceText) == 0:
                    ProjectExperienceText = "Geen ervaring beschikbaar"
        
        
        if isinstance(ProjectExperience,unicode) == True:
            print "type is string"
            ProjectExperienceText = ProjectExperience
            
        return ProjectExperienceText
    
    
    
    #Write document
    while (i < 10 ):

        #Profile
        print "PROFILE: " + str(i)
        objProfile  =  profile[i]
        #Elements of profile 
        FullName = getFullName(objProfile)[0]
        documentName = "CV_2015_" + FullName + ".docx"
        
        #Define Template
        document = Document()
        document.add_heading('CURRICULUM VITAE',3)
        document.add_paragraph()
        table = document.add_table(rows=7,cols=4)
        document.add_paragraph()
        document.add_heading('Korte omschrijving:',4)
       
               
        Summary = ""
        Summary = getBackgroundSummary(objProfile)
  
        # Full Name 
        setFullName(FullName)  
            
        #Adres/Postcode/GSM/Email/
        setAdress("Onbekend")  
        setPostcode("Onbekend")
        setGSM("Onbekend")
        setEmail("Onbekend")
           
        #Geboorte Datum/Geslacht/Natioanaliteit
        setGebDat("Onbekend")
        setGender("Onbekend")
        setNationality("Nederlandse")
         
        #Korte omschrijving   
        txtSummary = document.add_paragraph().add_run(setSummary(Summary))
        font = txtSummary.font 
        font.name = 'Arial'
        font.size = Pt(10)
        
        #Ervaring       
        if str("background_experience" in objProfile.keys()) == "True":
            print "background_experience: Aanwezig"
            document.add_heading("Ervaring:",4)
            txtBackgroundSummary = document.add_paragraph().add_run(setBackgroundExperience(getBackgroundExperience(objProfile)))
            font = txtBackgroundSummary.font 
            font.name = 'Arial'
            font.size = Pt(10)
        else:
            document.add_heading("Ervaring:",4)
            txtBackgroundSummary = document.add_paragraph().add_run("Geen achtergrond ervaring beschikbaar")
            font = txtBackgroundSummary.font 
            font.name = 'Arial'
            font.size = Pt(10)
            
            
        #Projecten       
        if str("background_projects" in objProfile.keys()) == "True":
            print "background_projects: Aanwezig"
            document.add_heading("Projecten:",4)
            txtProjectExperience = document.add_paragraph().add_run(setProjectExperience(getBackgroundProjects(objProfile)))
            font = txtProjectExperience.font 
            font.name = 'Arial'
            font.size = Pt(10)
        else:
            document.add_heading("Projecten:",4)
            txtBackgroundSummary = document.add_paragraph().add_run("Geen projecten beschikbaar")
            font = txtBackgroundSummary.font 
            font.name = 'Arial'
            font.size = Pt(10)
            
        
        print objProfile['skills_list']      




            
        document.save("[" +str(i) + "]" + documentName)
        print "------------------------------------- "
        print documentName + " ....SAVED"
        print "------------------------------------- " 
        i = i + 1
        

        